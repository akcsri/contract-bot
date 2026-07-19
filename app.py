import os
import io
import re
import time
import json
import logging
import datetime
import unicodedata

import requests
import fitz  # PyMuPDF
from docx import Document

from google import genai
from google.genai import types

from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler

# --- ログ設定 -------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

SLACK_BOT_TOKEN = os.environ["SLACK_BOT_TOKEN"]
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")

# このチャンネルの投稿のみ処理する(テスト/本番の切り分け)
CONTRACT_CHANNEL_ID = os.environ["CONTRACT_CHANNEL_ID"]

# --- freee 設定 -------------------------------------------------------
FREEE_CLIENT_ID = os.environ["FREEE_CLIENT_ID"]
FREEE_CLIENT_SECRET = os.environ["FREEE_CLIENT_SECRET"]
# 「申請者(金子さん)本人」としてのrefresh_token。承認経路(FREEE_APPROVAL_FLOW_ROUTE_ID)
# が金子さんが申請者であることを前提にしているため、証憑アップロードと
# 申請作成(create_nda_approval_request)は必ずこちらを使う。
FREEE_REFRESH_TOKEN = os.environ["FREEE_REFRESH_TOKEN"]
# (任意) 管理者権限を持つ別ユーザーのrefresh_token。sections/approval_flow_routes
# など、金子さんのユーザー権限では読み取れないマスタ系APIの参照専用に使う。
# 未設定の場合はFREEE_REFRESH_TOKEN(金子さん)にフォールバックする
# (これまで通りuser_do_not_have_permissionになる可能性がある)。
FREEE_ADMIN_REFRESH_TOKEN = os.environ.get("FREEE_ADMIN_REFRESH_TOKEN")
FREEE_COMPANY_ID = int(os.environ["FREEE_COMPANY_ID"])
FREEE_NDA_FORM_ID = int(os.environ.get("FREEE_NDA_FORM_ID", "87137"))
# 「NDA契約締結申請」フォームの「申請経路の選択」に対応する必須項目。
# NDA締結申請の経路(プロジェクトオーナー→リーガル(金子)→コーポレート(吉田))のID。
# 注意: 以前使っていた1431338は/api/1/approval_flow_routesで確認したところ
# 「存在しないか既に削除された申請経路」だった。GET /api/1/approval_flow_routes
# の一覧から、request_form_ids に87137(NDA申請フォーム)を含み、name="NDA締結申請"
# の現在有効なIDである1430722に更新した。
FREEE_APPROVAL_FLOW_ROUTE_ID = int(os.environ.get("FREEE_APPROVAL_FLOW_ROUTE_ID", "1430722"))
# 認可コード取得時に指定したコールバックURLと同じ値を指定する
# (ブラウザで手動取得した場合は "urn:ietf:wg:oauth:2.0:oob")
FREEE_REDIRECT_URI = os.environ.get("FREEE_REDIRECT_URI", "urn:ietf:wg:oauth:2.0:oob")

FREEE_API_BASE = "https://api.freee.co.jp"
FREEE_TOKEN_URL = "https://accounts.secure.freee.co.jp/public_api/token"

# --- Render API (ローテーションするrefresh_tokenを永続化するため) --------
# 未設定の場合は永続化をスキップする(従来通りメモリ内のみの保持になる)。
RENDER_API_KEY = os.environ.get("RENDER_API_KEY")
RENDER_SERVICE_ID = os.environ.get("RENDER_SERVICE_ID")
RENDER_API_BASE = "https://api.render.com/v1"

# 契約締結方法の選択肢(freee上のフォームの選択肢と一致させること。
# 増やす/変える場合はここを編集してください)
CONTRACT_METHODS = ["電子署名（CSRI発信）", "原本捺印"]

# freeeへの送信を承認したとみなすリアクション(そのまま申請する)
APPROVE_REACTIONS = {"+1", "thumbsup", "white_check_mark", "heavy_check_mark"}
# 下書きとして保存し、担当者が後で内容を確認・修正してから
# freee上で自分で「申請」を押すことを選んだとみなすリアクション
DRAFT_REACTIONS = {"memo", "pencil2", "spiral_note_pad"}

app = App(
    token=SLACK_BOT_TOKEN,
    logger=logger,
)

gemini_client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])

# 二重処理防止(Slackのretry/イベント再送対策の簡易ガード)
_processed_file_ids = set()

# NDA判定〜freee申請までの一時状態(thread_ts をキーに保持)
# 注意: プロセス内メモリのみ。Renderの再起動で消える簡易実装。
_pending_nda = {}

# --- プロジェクト名 → freee部門(section)ID の対応表(フォールバック用) ---
# 管理者権限トークンが使えるようになったことで、プロジェクト名の判定は
# 通常 fetch_freee_sections()/find_section_id_by_name() 経由でfreeeの
# 実データを直接参照する(この辞書には頼らない)。
# この辞書は、freeeのAPIが一時的に使えなくなった場合のフォールバックや、
# 表記ゆれのエイリアス登録(例: freee上の正式名称と少し違う呼び方をSlackで
# 使いたい場合)にのみ使う。
KNOWN_PROJECT_SECTION_IDS = {
    "CSRI": 3269199,
    "アバウテック": 3247109,
}

# --- 承認者名 → freeeユーザーID の対応表 -------------------------------
# 承認者は完全に人の判断による選択であり、プロジェクトや契約内容からは
# 自動的に決まらないため、Slackスレッドで都度名前を確認し、この対応表で
# freeeのユーザーIDに変換してから申請する。
#
# 新しい承認者のIDを調べる方法:
#   1. freeeにログインし、「設定」→「メンバー管理」で対象者のプロフィールを開く、
#      または対象者自身にfreeeの自分のユーザーIDを確認してもらう
#   2. もしくは実際にfreee上でその人を承認者に選んで「申請」を押す際の
#      ブラウザDevTools上のリクエストペイロードに approver_id として現れる値を使う
#   3. その<ID>と承認者名をこの辞書に追記する
#
# 注意: freeeにはHR(人事労務)側のemployee_idと会計側のuser_idという
# 別の採番体系が存在する(例: 金子明彦はemployee_id=3629707, user_id=13323251)。
# ここに登録するのは必ず会計側のuser_id(申請/承認APIで使われる方)。
KNOWN_APPROVERS = {
    "前田拓": 13236467,
    "堀内駿太郎": 13317152,
    "金井俊太朗": 14602813,
    "吉田愛美": 13233207,
    "塩谷輝": 13323276,
    "金子明彦": 13323251,
    "川端真至": 13323254,
    "嶋﨑江美": 13323278,
    "池内遼": 14387003,
    "村中康平": 14556353,
    "石谷虎太郎": 14556348,
}


def normalize_name(name: str) -> str:
    """全角/半角・空白種類の違いを吸収してから比較するための正規化。
    (Slackでの手入力やIMEの変換で、全角空白や異体の空白文字が
    混ざっても一致するようにする)
    """
    return unicodedata.normalize("NFKC", name).translate(
        {ord(c): None for c in " 　\t​"}
    )


_NORMALIZED_KNOWN_APPROVERS = {normalize_name(k): v for k, v in KNOWN_APPROVERS.items()}


def find_approver_id_by_name(name: str):
    return _NORMALIZED_KNOWN_APPROVERS.get(normalize_name(name))


# 金子明彦さんご本人のSlackユーザーID(識別用の特別扱いに使う)
AKIHIKO_SLACK_USER_ID = "U07PUQG9NNQ"

# --- Slackユーザー → freeeユーザーID(申請者)の対応表 --------------------
# 当初は「申請者(applicant_id)はAPIリクエストで独立に指定でき、金子さんの
# トークンのままapplicant_idだけ切り替えれば済むはず」と考えていたが、
# 実地検証(吉田さんの投稿でapplicant_id=吉田さんのIDを指定)の結果、
# レスポンスのapplicant_idは常に金子さんのままだった。つまり、
# 申請者は「実際にAPIを呼び出したfreeeアカウント本人」で決まり、
# request body上のapplicant_idでは上書きできないことが判明した。
#
# そのため、この辞書(Slack投稿者→freeeユーザーID)は主に確認メッセージの
# 表示用に使う。実際にその人本人の名義で申請するには、KNOWN_FREEE_USERSの
# 登録に加えて、その人自身のrefresh_tokenを環境変数
# FREEE_REFRESH_TOKEN_<SlackユーザーID> として登録する必要がある
# (resolve_freee_identity()参照)。
#
# 新しいSlackユーザーを追加する方法:
#   1. Slackプロフィールの「その他」→「メンバーIDをコピー」でSlackユーザーID
#      (例: U07PUQG9NNQ)を確認する
#   2. freeeにログインし、「設定」→「メンバー管理」で対象者のプロフィールを開き
#      ユーザーIDを確認する(またはKNOWN_APPROVERSと同様にDevTools等で調べる)
#   3. 下の辞書に {"<SlackユーザーID>": <freeeユーザーID>} を追記する
KNOWN_FREEE_USERS = {
    "U04LT7F25A7": 13236467,  # 前田拓
    "U04LVQSCWQL": 13317152,  # 堀内駿太郎
    "U05ECQYLYF4": 14602813,  # 金井俊太朗
    "U06MND3BB6E": 13233207,  # 吉田愛美
    "U07DBCX31EV": 13323276,  # 塩谷輝
    "U07PUQG9NNQ": 13323251,  # 金子明彦
    "U07UBPGDT7X": 13323254,  # 川端真至
    "U087HQL9RHC": 13323278,  # 嶋﨑江美
    "U09QGF7GB8E": 14387003,  # 池内遼
    "U09QC4PLNHZ": 14556353,  # 村中康平
    "U09R5UCKASF": 14556348,  # 石谷虎太郎
}


def find_applicant_id_by_slack_user(slack_user_id: str):
    return KNOWN_FREEE_USERS.get(slack_user_id)

PDF_MIMETYPES = {"application/pdf"}
PDF_FILETYPES = {"pdf"}
DOCX_MIMETYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DOCX_FILETYPES = {"docx"}

# 不足項目を聞き返すときに使う個別パターン(どちらか一方だけの返信でも拾える)
PROJECT_REPLY_PATTERN = re.compile(r"プロジェクト名?[:：]\s*(\S+)")
METHOD_REPLY_PATTERN = re.compile(r"締結方法[:：]\s*(.+)")
APPROVER_REPLY_PATTERN = re.compile(r"承認者[:：]\s*(\S+)")


def build_nda_prompt(project_candidates: list) -> str:
    candidates_text = "、".join(project_candidates) if project_candidates else "(候補リスト取得不可)"
    return f"""\
あなたは契約書の申請処理を行う専門アシスタントです。
提供された契約書本体(PDFの場合はスキャン画像のみでテキスト層が無い
こともあるので、画像として内容を読み取って判断してください)に加えて、
Slack投稿の本文とファイル名も参考情報として渡されます。契約書の内容
だけでなく、投稿本文・ファイル名に書かれている情報もすべて確認し、
そこに既に書かれている情報は憶測せずそのまま採用してください。
どこにも情報が無い項目だけをnullにしてください。

以下の項目をJSON形式のみで回答してください。説明文やコードブロックの
マークダウンは付けず、JSONオブジェクトのみを返すこと。

{{
  "is_nda": true または false (秘密保持契約/NDAであればtrue),
  "document_type": "書類の種類(例: 秘密保持契約書, 業務委託契約書 等)",
  "document_title": "契約書の冒頭(表紙・見出し部分)に実際に印字されている
    契約書名をそのままの文字列で抽出する(例: '守秘義務に関する確約書'。
    書類種別の一般的な分類ではなく、その書類に実際に書かれている
    タイトルの表記そのものを優先する)。読み取れない場合はnull",
  "parties": ["契約当事者1", "契約当事者2"],
  "counterparty_name": "契約当事者のうち、株式会社企業支援総合研究所(CSRI)
    ではない側(相手方)の正式な会社名または氏名を、契約書に記載されている
    通りの表記で1つだけ抽出する。当事者名にプレースホルダー(●●、対象会社
    など)しか無く実名が分からない場合や、相手方が判別できない場合はnull",
  "contract_date": "YYYY-MM-DD形式の契約日。読み取れない場合はnull",
  "reason": "is_ndaと判定した理由の要約(1〜2文)",
  "project_name": "次の候補の中から一致するものを1つだけ選んで文字列で返す: [{candidates_text}]。
    ファイル名や投稿本文にプロジェクト名/顧客名らしき記載があれば最優先で使い、
    候補の中から最も一致するものを選ぶこと。プロジェクトに紐づかない場合や
    候補に一致するものが無い場合はnull",
  "method": "契約の締結方法。投稿本文に「捺印」「押印」とあれば\\"原本捺印\\"、
    「電子署名」とあれば\\"電子署名（CSRI発信）\\"。判断できなければnull",
  "physical_mail_address": "締結方法が原本捺印と判断できる場合、投稿本文に
    書かれている原本の送付先(郵便番号・住所・会社名・部署名・担当者名・
    電話番号など)をそのままの文字列でまとめて抽出する。記載が無い/
    該当しない場合はnull"
}}
"""


# =====================================================================
# ファイル判定・ダウンロード
# =====================================================================

def is_pdf_file(f: dict) -> bool:
    return f.get("mimetype") in PDF_MIMETYPES or f.get("filetype") in PDF_FILETYPES


def is_docx_file(f: dict) -> bool:
    return f.get("mimetype") in DOCX_MIMETYPES or f.get("filetype") in DOCX_FILETYPES


def download_slack_file(file_info: dict) -> bytes:
    """Slackのfiles:read権限でファイル本体(バイナリ)を取得する"""
    url = file_info["url_private"]
    resp = requests.get(
        url,
        headers={"Authorization": f"Bearer {SLACK_BOT_TOKEN}"},
        timeout=30,
    )
    resp.raise_for_status()
    return resp.content


def get_pdf_page_count(pdf_bytes: bytes) -> int:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        return doc.page_count
    finally:
        doc.close()


def extract_docx_text(docx_bytes: bytes) -> str:
    """python-docxで段落・表のテキストを抽出する"""
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


# =====================================================================
# Gemini解析
# =====================================================================

def analyze_contract_with_gemini(
    *, pdf_bytes: bytes = None, text: str = None,
    message_text: str = "", filename: str = "",
    project_candidates: list = None,
) -> dict:
    """契約書(PDFバイナリ、またはWordから抽出済みのテキスト)に加えて、
    Slack投稿本文・ファイル名・freeeのプロジェクト候補一覧もあわせてGeminiに渡し、
    NDA判定と、freee申請に必要な項目(プロジェクト名/締結方法/原本送付先)の
    自動抽出を1回のリクエストでまとめて行う。
    PDFはスキャン画像(テキスト層無し)でもGeminiが画像として読み取れるため
    OCR処理を別途行う必要がない。Wordは事前にテキスト抽出したものを渡す。
    """
    prompt = build_nda_prompt(project_candidates or [])
    extra_context = (
        f"--- Slack投稿本文 ---\n{message_text or '(本文なし)'}\n\n"
        f"--- 添付ファイル名 ---\n{filename}\n"
    )

    if pdf_bytes is not None:
        contents = [
            types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
            prompt,
            extra_context,
        ]
    elif text is not None:
        contents = [
            prompt,
            extra_context,
            f"--- 契約書本文(Wordファイルから抽出) ---\n{text}",
        ]
    else:
        raise ValueError("pdf_bytes または text のいずれかが必要です")

    response = gemini_client.models.generate_content(
        model=GEMINI_MODEL,
        contents=contents,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )
    return json.loads(response.text)


def format_parties(result: dict) -> str:
    """Geminiが返すpartiesにNoneや空文字が混ざることがあるため除去してから結合する"""
    parties = [p for p in (result.get("parties") or []) if p]
    return "、".join(parties) or "不明"


# 自社名(取引先扱いにしないためのフィルタ)
OWN_COMPANY_NAMES = {"株式会社企業支援総合研究所", "CSRI"}


def resolve_counterparty_name(result: dict) -> str:
    """相手方の名称を決める。Geminiがcounterparty_nameを特定できなかった場合
    (プレースホルダーの契約書ひな形など)は、partiesのうち自社以外の
    記載をフォールバックとして使う(取引先欄を空のまま申請すると
    「取引先が入力されていません」というエラーになるため、名前が
    プレースホルダーだとしても何かしら入れておく)。
    """
    name = (result.get("counterparty_name") or "").strip()
    if name:
        return name
    parties = [p for p in (result.get("parties") or []) if p and p not in OWN_COMPANY_NAMES]
    return parties[0] if parties else ""


def format_analysis_message(filename: str, meta_line: str, result: dict) -> str:
    is_nda = result.get("is_nda")
    judgement = "✅ NDA(秘密保持契約)と判定" if is_nda else "❌ NDAではないと判定"
    parties = format_parties(result)
    return (
        f"ファイルを受信しました: *{filename}* ({meta_line})\n"
        f"{judgement}\n"
        f"書類種別: {result.get('document_type', '不明')}\n"
        f"契約当事者: {parties}\n"
        f"契約日: {result.get('contract_date') or '不明'}\n"
        f"判定理由: {result.get('reason', '-')}"
    )


# =====================================================================
# freee 連携
# =====================================================================

# freeeの「申請者」はAPIを実際に呼び出したアカウント本人になり、
# request body上のapplicant_idでは上書きできないことを実地検証済み
# (吉田さんのSlack投稿で、applicant_id=吉田さんのIDを指定しても
# レスポンスは金子さんのIDになった)。そのため、Slack投稿者本人の名義で
# 申請するには、その人自身のrefresh_tokenでAPIを呼ぶ必要がある。
#
# アイデンティティ(=それぞれ別のrefresh_token)を複数管理する:
#   "user"  : 金子さん本人。デフォルト(登録が無い投稿者はここにフォールバック)。
#   "admin" : 管理者権限を持つユーザー(吉田さん)。sections/approval_flow_routes
#             など、金子さんの権限では読めないマスタ系APIの参照用に加え、
#             吉田さんご本人が投稿した場合の申請者トークンとしても使う。
#   "<SlackユーザーID>" : その他の個人。環境変数
#             FREEE_REFRESH_TOKEN_<SlackユーザーID> (例:
#             FREEE_REFRESH_TOKEN_U07UBPGDT7X) が設定されていれば、
#             起動時に自動でここに登録される。追加のコード変更は不要で、
#             Renderに環境変数を追加するだけでよい。
_freee_token_caches = {
    "user": {
        "refresh_token": FREEE_REFRESH_TOKEN,
        "render_env_key": "FREEE_REFRESH_TOKEN",
        "access_token": None,
        "expires_at": 0,
    },
    "admin": {
        "refresh_token": FREEE_ADMIN_REFRESH_TOKEN or FREEE_REFRESH_TOKEN,
        "render_env_key": "FREEE_ADMIN_REFRESH_TOKEN" if FREEE_ADMIN_REFRESH_TOKEN else "FREEE_REFRESH_TOKEN",
        "access_token": None,
        "expires_at": 0,
    },
}

# 環境変数 FREEE_REFRESH_TOKEN_<SlackユーザーID> を自動的にスキャンして
# 個人アイデンティティとして登録する。
_PERSONAL_TOKEN_ENV_PREFIX = "FREEE_REFRESH_TOKEN_"
for _env_key, _env_value in os.environ.items():
    if _env_key.startswith(_PERSONAL_TOKEN_ENV_PREFIX) and _env_value:
        _slack_id = _env_key[len(_PERSONAL_TOKEN_ENV_PREFIX):]
        _freee_token_caches[_slack_id] = {
            "refresh_token": _env_value,
            "render_env_key": _env_key,
            "access_token": None,
            "expires_at": 0,
        }
        logger.info(f"[freee] 個人アイデンティティを登録しました: {_slack_id} (env: {_env_key})")

# Slackユーザー→識別子の手動マッピング(環境変数の命名規則に頼らない特例用)。
# 吉田さんは"admin"アイデンティティ(FREEE_ADMIN_REFRESH_TOKEN)を
# 申請者トークンとしてもそのまま流用する。
SLACK_USER_TO_FREEE_IDENTITY = {
    "U06MND3BB6E": "admin",  # 吉田愛美
}


def resolve_freee_identity(slack_user_id: str) -> str:
    """Slack投稿者のユーザーIDから、その人自身のfreeeトークンのアイデンティティ
    キーを解決する。個別トークンが登録されていなければ"user"(金子さん)
    にフォールバックする。
    """
    if slack_user_id in SLACK_USER_TO_FREEE_IDENTITY:
        return SLACK_USER_TO_FREEE_IDENTITY[slack_user_id]
    if slack_user_id in _freee_token_caches:
        return slack_user_id
    return "user"


def _describe_applicant_identity(slack_user_id: str) -> str:
    """確認メッセージ用に、申請者名義の状態を説明する文言を返す。"""
    identity = resolve_freee_identity(slack_user_id)
    if identity != "user":
        return "(投稿者ご本人名義で申請されます)"
    if slack_user_id == AKIHIKO_SLACK_USER_ID:
        return "(ご本人名義で申請されます)"
    return "(投稿者ご本人のトークンが未登録のため、金子明彦名義で申請されます)"


_sections_cache = None


def get_freee_access_token(identity: str = "user") -> str:
    """freeeのアクセストークンを取得する(必要な時だけrefresh)。identityで
    "user"(金子さん本人)か"admin"(管理者)のどちらのトークンを使うか選ぶ。

    重要な注意点: freeeのrefresh_tokenは使用するたびに新しい値に
    ローテーション(再発行)される。このプロセスはローテーション後の
    refresh_tokenをメモリ内にしか保持していないため、Renderのプロセスが
    再起動すると環境変数が指す値がすでに無効になっている可能性がある。
    そのためローテーション後の値は毎回Renderの環境変数に書き戻している
    (persist_refresh_token_to_render)。
    """
    cache = _freee_token_caches[identity]
    now = time.time()
    if cache["access_token"] and now < cache["expires_at"] - 60:
        return cache["access_token"]

    logger.info(
        f"[freee:{identity}] refreshing access token (redirect_uri={FREEE_REDIRECT_URI}, "
        f"refresh_token末尾={cache['refresh_token'][-6:]})"
    )
    resp = requests.post(
        FREEE_TOKEN_URL,
        data={
            "grant_type": "refresh_token",
            "client_id": FREEE_CLIENT_ID,
            "client_secret": FREEE_CLIENT_SECRET,
            "refresh_token": cache["refresh_token"],
            "redirect_uri": FREEE_REDIRECT_URI,
        },
        timeout=30,
    )
    if not resp.ok:
        logger.error(f"[freee:{identity}] token refresh failed: {resp.status_code} {resp.text}")
    resp.raise_for_status()
    data = resp.json()
    cache["access_token"] = data["access_token"]
    cache["refresh_token"] = data["refresh_token"]
    cache["expires_at"] = now + data["expires_in"]
    logger.info(f"[freee:{identity}] access token refreshed")
    persist_refresh_token_to_render(cache["render_env_key"], data["refresh_token"])
    return cache["access_token"]


def persist_refresh_token_to_render(env_key: str, new_refresh_token: str):
    """ローテーションされたrefresh_tokenをRenderの環境変数(env_key)に書き戻す。

    これをしないと、プロセス再起動のたびに環境変数の古い(すでに使用済みの)
    refresh_tokenが読み込まれ、invalid_grantで失敗し続ける。
    RENDER_API_KEY / RENDER_SERVICE_ID が未設定の場合、またはenv_keyに対応する
    環境変数がそもそも設定されていない場合(admin用トークン未使用時)は何もしない。

    注意: Renderの環境変数を更新すると、そのサービスは自動的に再デプロイ
    される。そのため、アクセストークンの更新(=このタイミング)のたびに
    短い再起動が発生する。
    """
    if not RENDER_API_KEY or not RENDER_SERVICE_ID:
        logger.warning(
            f"[render] RENDER_API_KEY/RENDER_SERVICE_ID未設定のため、"
            f"{env_key}の永続化をスキップします(再起動すると失効する可能性があります)"
        )
        return
    try:
        resp = requests.put(
            f"{RENDER_API_BASE}/services/{RENDER_SERVICE_ID}/env-vars/{env_key}",
            headers={
                "Authorization": f"Bearer {RENDER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={"value": new_refresh_token},
            timeout=15,
        )
        if not resp.ok:
            logger.error(f"[render] env var update failed: {resp.status_code} {resp.text}")
            return
        logger.info(f"[render] {env_key}をRenderに永続化しました")
    except Exception:
        logger.exception(f"[render] {env_key}の永続化に失敗しました")


def freee_headers(identity: str = "user") -> dict:
    return {"Authorization": f"Bearer {get_freee_access_token(identity)}"}


def fetch_freee_sections() -> list:
    """freeeの「部門」マスタをAPIから取得する(プロセス内でキャッシュする)。

    以前は金子さん本人のユーザー権限ではuser_do_not_have_permissionで
    使えなかったが、管理者権限のトークン("admin"アイデンティティ)を
    使うようになったことで取得できるようになった。これにより、
    プロジェクト名の判定はKNOWN_PROJECT_SECTION_IDS(手動辞書)に頼らず、
    毎回freeeの最新データを直接参照できる。
    """
    global _sections_cache
    if _sections_cache is not None:
        return _sections_cache
    resp = requests.get(
        f"{FREEE_API_BASE}/api/1/sections",
        headers=freee_headers("admin"),
        params={"company_id": FREEE_COMPANY_ID},
        timeout=30,
    )
    if not resp.ok:
        logger.error(f"[freee] sections fetch failed: {resp.status_code} {resp.text}")
    resp.raise_for_status()
    _sections_cache = resp.json().get("sections", [])
    return _sections_cache


def get_project_candidates() -> list:
    """Geminiのプロジェクト名推定に使う候補名一覧。

    まずfreeeから実際に登録されているプロジェクト(部門)名を取得して使う。
    API呼び出しに失敗した場合(権限が再び失われた等)のみ、
    手動管理のKNOWN_PROJECT_SECTION_IDSにフォールバックする。
    """
    try:
        sections = fetch_freee_sections()
        names = [s.get("name") for s in sections if s.get("name")]
        if names:
            return names
    except Exception:
        logger.exception("[freee] sections取得に失敗したため、手動辞書にフォールバックします")
    return list(KNOWN_PROJECT_SECTION_IDS.keys())


def find_section_id_by_name(name: str):
    """プロジェクト名(freee上は部門名)からIDを引く。

    まずfreeeの実データ(fetch_freee_sections)から完全一致で探し、
    見つからなければ手動辞書(KNOWN_PROJECT_SECTION_IDS。表記ゆれの
    エイリアスなどを追加したい場合用)にフォールバックする。
    """
    if not name:
        return None
    try:
        sections = fetch_freee_sections()
        for s in sections:
            if s.get("name") == name:
                return s.get("id")
    except Exception:
        logger.exception("[freee] sections取得に失敗したため、手動辞書にフォールバックします")
    return KNOWN_PROJECT_SECTION_IDS.get(name)


def find_partner_id_by_name(name: str):
    """相手方の会社名/氏名から、既存のfreee取引先(partners)を検索してIDを返す。

    現在のアプリのOAuthスコープにaccounting:partners:readが含まれていない
    可能性があり、その場合は403/401になる。その場合は例外を握りつぶして
    Noneを返し、呼び出し元で「会社名をそのまま記入する」フォールバックを
    使う(取引先マスタへの読み取り権限が付与されれば自動的に使われるように
    しておく)。
    """
    if not name:
        return None
    try:
        resp = requests.get(
            f"{FREEE_API_BASE}/api/1/partners",
            headers=freee_headers("user"),
            params={"company_id": FREEE_COMPANY_ID, "keyword": name},
            timeout=15,
        )
        if not resp.ok:
            logger.warning(
                f"[freee] 取引先検索に失敗(スコープ不足の可能性、名前で代用します): "
                f"{resp.status_code} {resp.text}"
            )
            return None
        partners = resp.json().get("partners", [])
        for p in partners:
            if p.get("name") == name:
                return p.get("id")
        for p in partners:
            pname = p.get("name") or ""
            if name in pname or pname in name:
                return p.get("id")
        return None
    except Exception:
        logger.exception("[freee] 取引先検索でエラーが発生しました(名前で代用します)")
        return None


def resolve_counterparty_fields(counterparty_name: str) -> tuple:
    """相手方に関する2つの欄の値を決める。

    フォーム定義(GET /api/1/approval_requests/forms/87137)によると、
    id 32070(type: partner, ラベル「契約相手方（既存取引先）」)と
    id 789358(type: single_line, ラベル「契約相手方（新規取引先）」)は、
    どちらか一方に入力する二者択一の項目(両方空だと
    「取引先が入力されていません」というエラーになることを確認済み)。
    既存の取引先に一致するものがあればそのIDをpartner欄に、
    無ければ会社名/氏名の文字列をsingle_line(新規取引先)欄に入れる。

    戻り値: (partner_value, single_line_value)
    """
    if not counterparty_name:
        return "", ""
    partner_id = find_partner_id_by_name(counterparty_name)
    if partner_id is not None:
        logger.info(f"[freee] 取引先「{counterparty_name}」を既存の取引先ID {partner_id} に解決しました")
        return str(partner_id), ""
    logger.info(f"[freee] 取引先「{counterparty_name}」に一致する既存取引先が無いため、新規取引先名として使用します")
    return "", counterparty_name


def upload_file_to_freee(file_bytes: bytes, filename: str, identity: str = "user") -> int:
    """freeeのファイルボックス(証憑)にアップロードし、receipt idを返す。

    identityには、Slack投稿者本人のfreeeトークンを使うべきかどうかを
    resolve_freee_identity()で解決した値を渡す(申請者=証憑をアップロード
    したアカウント、という前提と揃えるため)。
    """
    resp = requests.post(
        f"{FREEE_API_BASE}/api/1/receipts",
        headers=freee_headers(identity),
        data={"company_id": FREEE_COMPANY_ID, "description": filename},
        files={"receipt": (filename, file_bytes)},
        timeout=60,
    )
    if not resp.ok:
        logger.error(f"[freee] receipt upload failed: {resp.status_code} {resp.text}")
    resp.raise_for_status()
    return resp.json()["receipt"]["id"]


def create_nda_approval_request(
    *, title: str, contract_date: str,
    receipt_id: int, section_id: int, method: str, approver_id: int = None,
    applicant_id: int = None, partner_value: str = "", new_partner_name: str = "",
    mail_address: str = "", draft: bool = False, identity: str = "user",
) -> dict:
    """NDA契約締結申請(freeeの汎用申請フォーム, form_id=87137)を作成する。

    各項目には、freee側のフォーム定義に紐づく固定の`id`を指定する必要がある
    (`type`と`value`だけでは「Idを入力してください」というバリデーション
    エラーになる)。このidはブラウザの実際のフォーム(下書き保存時のリクエスト
    ペイロード)から採取したもので、フォーム定義が変わらない限り固定。
    フォームの項目が追加/削除/並び替えされた場合はここも合わせて調整すること。
    最初のmulti_lineは「原本送付先」欄(押印方法が原本捺印の場合のみ使用)。

    draft=Falseの場合(そのまま申請): 承認者(approver_id)は必須。
    プロジェクトや契約内容から自動的に決まるものではなく完全に人の判断
    による選択のため、Slackスレッドで都度確認し(KNOWN_APPROVERSで
    名前→IDに変換したうえで)呼び出し元から渡す。nullを指定すると
    「approver_idはnullを指定することはできません」というエラーになる。

    draft=Trueの場合(下書き保存): 承認者はまだ未定でよいため、
    approver_idキー自体をリクエストボディから省略する(担当者が後で
    freee上で承認者を選んで「申請」を押す想定)。

    申請者は、request body上のapplicant_idでは上書きできず、実際に
    このAPIを呼び出したfreeeアカウント本人になることを実地検証済み
    (吉田さんの投稿でapplicant_idに吉田さんのIDを指定しても、
    レスポンスは金子さんのIDになった)。そのため、投稿者本人名義で
    申請するには、identity引数でその人自身のfreeeトークンを指定して
    呼び出す必要がある(resolve_freee_identity()で解決する)。
    applicant_idはあくまで参考情報として送っているが、実際の申請者を
    決めるのはidentityの方である点に注意。
    """
    body = {
        "company_id": FREEE_COMPANY_ID,
        "form_id": FREEE_NDA_FORM_ID,
        "approval_flow_route_id": FREEE_APPROVAL_FLOW_ROUTE_ID,
        "title": title,
        "draft": draft,
        "applicant_id": applicant_id,
        # group_id/applicant_group_id/observer_user_idsはnullでも明示的に含めないと
        # approval_flow_route_idが不正というエラーになることを確認済み。
        "group_id": None,
        "applicant_group_id": None,
        "observer_user_ids": [],
        "request_items": [
            {"id": 346116, "type": "title", "value": title},
            {"id": 57980, "type": "section", "value": str(section_id)},
            {"id": 789358, "type": "single_line", "value": new_partner_name},
            {"id": 32070, "type": "partner", "value": partner_value},
            {"id": 293716, "type": "date", "value": contract_date},
            {"id": 555747, "type": "receipt", "value": str(receipt_id)},
            {"id": 647567, "type": "select", "value": method},
            {"id": 757585, "type": "multi_line", "value": mail_address},
            {"id": 757586, "type": "multi_line", "value": ""},
            {"id": 757587, "type": "multi_line", "value": ""},
        ],
    }
    if not draft:
        body["approver_id"] = approver_id
    resp = requests.post(
        f"{FREEE_API_BASE}/api/1/approval_requests",
        headers=freee_headers(identity),
        json=body,
        timeout=30,
    )
    if not resp.ok:
        logger.error(f"[freee] approval request creation failed: {resp.status_code} {resp.text}")
    resp.raise_for_status()
    return resp.json()["approval_request"]


# =====================================================================
# Slack: 契約書ファイル受信 → Gemini解析 → NDA判定 → freee申請準備
# =====================================================================

def post_confirmation(pending: dict, thread_ts: str, say):
    """freee申請の最終確認メッセージを投稿する"""
    result = pending["gemini_result"]
    contract_title = (
        (result.get("document_title") or "").strip()
        or (result.get("document_type") or "").strip()
        or os.path.splitext(pending["filename"])[0]
    )
    counterparty_name = resolve_counterparty_name(result) or "不明"
    lines = [
        "以下の内容でfreeeへの登録を行います。",
        "・そのまま申請する場合: このメッセージに :+1: で反応してください",
        "・内容を後で担当者が確認・修正してから申請する場合"
        "(下書き保存のみ): このメッセージに :memo: で反応してください",
        f"タイトル: {contract_title}",
        f"契約当事者: {format_parties(result)}",
        f"相手方(取引先欄): {counterparty_name}",
        f"契約日: {result.get('contract_date') or '不明'}",
        f"プロジェクト名: {pending['section_name']}"
        + ("(特定できなかったため自動設定。違う場合はfreee上で修正してください)"
           if pending.get("project_auto_defaulted") else ""),
        f"締結方法: {pending['method']}",
        f"承認者: {pending['approver_name']}",
        f"申請者: freeeユーザーID {pending['applicant_id']}"
        + _describe_applicant_identity(pending.get("posting_slack_user_id")),
    ]
    if pending.get("mail_address"):
        lines.append(f"原本送付先: {pending['mail_address']}")
    posted = say("\n".join(lines), thread_ts=thread_ts)
    pending["confirm_message_ts"] = posted["ts"]
    pending["stage"] = "awaiting_confirm"


def prompt_for_missing_fields(missing_fields: list, thread_ts: str, say):
    prompts = []
    if "project_name" in missing_fields:
        prompts.append("プロジェクト名: <プロジェクト名>")
    if "method" in missing_fields:
        prompts.append(f"締結方法: <{' か '.join(CONTRACT_METHODS)}>")
    if "approver" in missing_fields:
        prompts.append(f"承認者: <{' か '.join(KNOWN_APPROVERS.keys())}>")
    note = (
        "\n(プロジェクトに紐づかない契約書は「CSRI」を指定してください)"
        if "project_name" in missing_fields else ""
    )
    say(
        "投稿内容・ファイル名・契約書の中身から自動判定を試みましたが、"
        "次の項目が確認できませんでした。このスレッドで返信してください"
        "(承認者は毎回、人が選んで確認する運用のため必ず聞いています)。\n"
        f"`{' / '.join(prompts)}`" + note,
        thread_ts=thread_ts,
    )


def handle_contract_files(event: dict, say):
    files = event.get("files", [])
    target_files = [f for f in files if is_pdf_file(f) or is_docx_file(f)]

    if not target_files:
        return

    message_ts = event.get("ts")
    message_text = event.get("text", "")
    posting_slack_user_id = event.get("user")

    # プロジェクト名の自動判定に使う候補一覧(KNOWN_PROJECT_SECTION_IDSより)
    project_candidates = get_project_candidates()

    for f in target_files:
        file_id = f.get("id")
        filename = f.get("name", "unknown")

        if file_id in _processed_file_ids:
            logger.info(f"[SKIP] already processed: {filename} ({file_id})")
            continue
        _processed_file_ids.add(file_id)

        try:
            logger.info(f"[FILE] downloading: {filename} ({file_id})")
            raw_bytes = download_slack_file(f)

            if is_pdf_file(f):
                page_count = get_pdf_page_count(raw_bytes)
                logger.info(f"[PDF] sending to Gemini({GEMINI_MODEL}): {filename}")
                result = analyze_contract_with_gemini(
                    pdf_bytes=raw_bytes,
                    message_text=message_text,
                    filename=filename,
                    project_candidates=project_candidates,
                )
                meta_line = f"ページ数: {page_count}"

            else:  # docx
                text = extract_docx_text(raw_bytes)
                if not text.strip():
                    logger.warning(f"[DOCX] no extractable text: {filename}")
                    say(f":warning: Wordファイルからテキストを抽出できませんでした: {filename}")
                    continue
                logger.info(f"[DOCX] sending to Gemini({GEMINI_MODEL}): {filename}")
                result = analyze_contract_with_gemini(
                    text=text,
                    message_text=message_text,
                    filename=filename,
                    project_candidates=project_candidates,
                )
                meta_line = f"文字数: {len(text)}"

            logger.info(f"[FILE] gemini result: {result}")
            say(format_analysis_message(filename, meta_line, result))

            if not result.get("is_nda"):
                continue

            thread_ts = message_ts

            project_name = (result.get("project_name") or "").strip()
            method = (result.get("method") or "").strip()
            mail_address = (result.get("physical_mail_address") or "").strip()

            section_id = None
            if project_name:
                section_id = find_section_id_by_name(project_name)

            project_auto_defaulted = False
            if section_id is None:
                # 自信をもって特定できなかった場合は、プロジェクトに
                # 紐づかない契約として扱い、自動でCSRIにフォールバックする
                # (Slackでの聞き返しはしない。必要ならfreee上で本人が修正する)。
                fallback_id = find_section_id_by_name("CSRI")
                if fallback_id is not None:
                    section_id = fallback_id
                    project_name = "CSRI"
                    project_auto_defaulted = True

            missing_fields = []
            if section_id is None:
                # CSRI自体が対応表に無い場合のみ(通常起きない)、聞き返しにフォールバック
                missing_fields.append("project_name")
                project_name = ""
            if method not in CONTRACT_METHODS:
                missing_fields.append("method")
                method = ""
            # 承認者はプロジェクト/契約内容から自動的に決まらない人の判断のため、
            # 毎回Slackスレッドで確認する(自動判定・自動デフォルトはしない)。
            missing_fields.append("approver")

            # 申請者(applicant_id)はSlack投稿者に対応するfreeeユーザーID。
            # 未登録の場合は金子さん名義にフォールバックし、確認メッセージで
            # その旨を明示する(申請自体は失敗させない)。
            applicant_id = find_applicant_id_by_slack_user(posting_slack_user_id)
            applicant_auto_defaulted = applicant_id is None
            if applicant_id is None:
                applicant_id = KNOWN_FREEE_USERS.get(AKIHIKO_SLACK_USER_ID)  # 金子明彦にフォールバック

            pending = {
                "filename": filename,
                "raw_bytes": raw_bytes,
                "gemini_result": result,
                "section_id": section_id,
                "section_name": project_name,
                "project_auto_defaulted": project_auto_defaulted,
                "method": method,
                "mail_address": mail_address if method == "原本捺印" else "",
                "approver_id": None,
                "approver_name": "",
                "posting_slack_user_id": posting_slack_user_id,
                "applicant_id": applicant_id,
                "applicant_auto_defaulted": applicant_auto_defaulted,
                "missing_fields": missing_fields,
                "stage": "awaiting_fields" if missing_fields else "awaiting_confirm",
            }
            _pending_nda[thread_ts] = pending

            if missing_fields:
                prompt_for_missing_fields(missing_fields, thread_ts, say)
            else:
                post_confirmation(pending, thread_ts, say)

        except requests.HTTPError as e:
            logger.exception(f"[FILE] download failed: {filename}")
            say(f":warning: ファイルのダウンロードに失敗しました: {filename} ({e})")
        except json.JSONDecodeError as e:
            logger.exception(f"[FILE] gemini response was not valid JSON: {filename}")
            say(f":warning: Geminiの解析結果を読み取れませんでした: {filename} ({e})")
        except Exception as e:
            logger.exception(f"[FILE] processing failed: {filename}")
            say(f":warning: ファイルの処理に失敗しました: {filename} ({e})")


def strip_reply_chrome(s: str) -> str:
    """返信からヒント文をそのままコピペした際に混入しがちな記号を除去する
    (バッククォート `` ` `` や、ヒントの `<name>` 表記の山括弧など)。
    """
    return s.strip().strip("`<>「」\"'")


def handle_nda_field_reply(event: dict, say) -> bool:
    """NDA申請で不足している項目(プロジェクト名/締結方法)の返信を処理する。
    片方だけの返信でも受け付け、揃うまで不足分だけを聞き返す。
    処理した場合True、対象外ならFalseを返す。
    """
    thread_ts = event.get("thread_ts")
    pending = _pending_nda.get(thread_ts)
    if not pending or pending.get("stage") != "awaiting_fields":
        return False

    text = event.get("text", "")
    missing = list(pending.get("missing_fields", []))

    if "project_name" in missing:
        m = PROJECT_REPLY_PATTERN.search(text)
        if m:
            section_name = strip_reply_chrome(m.group(1))
            section_id = find_section_id_by_name(section_name)

            if section_id is None:
                say(
                    f":warning: プロジェクト「{section_name}」は対応表(KNOWN_PROJECT_SECTION_IDS)に"
                    "登録されていません。登録済みのプロジェクト名と完全に一致させて返信するか"
                    "(紐づかない場合は「CSRI」)、freee上でtags_history経由でIDを調べて"
                    "コードに追記してください。",
                    thread_ts=thread_ts,
                )
                return True

            pending["section_id"] = section_id
            pending["section_name"] = section_name
            missing.remove("project_name")

    if "method" in missing:
        m = METHOD_REPLY_PATTERN.search(text)
        if m:
            method = strip_reply_chrome(m.group(1))
            if method not in CONTRACT_METHODS:
                say(
                    f":warning: 締結方法は次のいずれかで指定してください: {', '.join(CONTRACT_METHODS)}",
                    thread_ts=thread_ts,
                )
                return True
            pending["method"] = method
            missing.remove("method")

    if "approver" in missing:
        m = APPROVER_REPLY_PATTERN.search(text)
        if m:
            approver_name = strip_reply_chrome(m.group(1))
            approver_id = find_approver_id_by_name(approver_name)
            logger.info(
                f"[approver] 抽出した名前: {approver_name!r} "
                f"/ 対応表のキー一覧: {[k for k in KNOWN_APPROVERS]!r} "
                f"/ 一致: {approver_id is not None}"
            )

            if approver_id is None:
                say(
                    f":warning: 承認者「{approver_name}」は対応表(KNOWN_APPROVERS)に"
                    "登録されていません。登録済みの名前と完全に一致させて返信するか、"
                    "freeeの承認者選択画面等でその人のユーザーIDを確認してコードに"
                    "追記してください。",
                    thread_ts=thread_ts,
                )
                return True

            pending["approver_id"] = approver_id
            pending["approver_name"] = approver_name
            missing.remove("approver")

    pending["missing_fields"] = missing

    if missing:
        prompt_for_missing_fields(missing, thread_ts, say)
        return True

    post_confirmation(pending, thread_ts, say)
    return True


@app.event("message")
def handle_message(event, say, logger):
    logger.info("========== EVENT RECEIVED ==========")
    logger.debug(json.dumps(event, ensure_ascii=False, indent=2))

    subtype = event.get("subtype")

    # bot自身の投稿は無視(無限ループ防止)
    if event.get("bot_id"):
        return

    # 指定チャンネル以外は無視(テスト/本番の切り分け)
    if event.get("channel") != CONTRACT_CHANNEL_ID:
        return

    # 通常のテキストメッセージ or ファイル添付(file_share)のみ処理
    # それ以外(message_changed / message_deleted 等)は無視
    if subtype not in (None, "file_share"):
        return

    if event.get("files"):
        handle_contract_files(event, say)
        return

    if event.get("thread_ts") and handle_nda_field_reply(event, say):
        return

    # 動作確認用コマンド: 権限が変わったfreeeトークンでsections/approval_flow_routes
    # が読めるようになったかを直接確認する(下書きコード内のfetch_freee_sectionsを使用)。
    # コピペ時にバッククォート等が混入しても一致するよう正規化してから比較する。
    text = strip_reply_chrome(event.get("text") or "")
    if text == "!debug_sections":
        try:
            sections = fetch_freee_sections()
            pairs = [f"{s.get('name')}={s.get('id')}" for s in sections]
            say(f":mag: sections取得成功({len(sections)}件): {pairs}")
        except Exception as e:
            say(f":warning: sections取得失敗: {e}")
        return

    # プロジェクト名(部門名)をキーワードで検索してIDを調べるコマンド。
    # 例: "!find_project Jaguar" と投稿すると、名前に"Jaguar"を含む
    # sectionをid付きで返す。KNOWN_PROJECT_SECTION_IDSに追記する際に使う。
    if text.startswith("!find_project"):
        keyword = strip_reply_chrome(text[len("!find_project"):])
        try:
            sections = fetch_freee_sections()
            matches = [
                f"{s.get('name')}={s.get('id')}"
                for s in sections
                if keyword.lower() in (s.get("name") or "").lower()
            ]
            if matches:
                say(f":mag: 「{keyword}」に一致するプロジェクト: {matches}")
            else:
                say(f":warning: 「{keyword}」に一致するプロジェクトが見つかりませんでした")
        except Exception as e:
            say(f":warning: プロジェクト検索失敗: {e}")
        return

    # 動作確認用コマンド: NDA申請フォーム(87137)自体の定義を取得する。
    # 「利用できない申請経路IDが指定されています」の原因調査用。
    # accounting:approval_requests:read スコープは既に付与済みのはずなので、
    # 金子さん本人のトークン("user")でアクセスできるはず。
    if text == "!debug_form":
        try:
            resp = requests.get(
                f"{FREEE_API_BASE}/api/1/approval_requests/forms/{FREEE_NDA_FORM_ID}",
                headers=freee_headers("user"),
                params={"company_id": FREEE_COMPANY_ID},
                timeout=30,
            )
            if not resp.ok:
                say(f":warning: フォーム定義取得失敗: {resp.status_code} {resp.text}")
            else:
                say(f":mag: フォーム定義: {json.dumps(resp.json(), ensure_ascii=False)[:3000]}")
        except Exception as e:
            say(f":warning: フォーム定義取得失敗: {e}")
        return

    # 動作確認用コマンド: 承認経路(1431338)自体の定義を取得する
    # (どの承認者/グループが正しいステップとして期待されているかを確認する)。
    if text == "!debug_route":
        try:
            resp = requests.get(
                f"{FREEE_API_BASE}/api/1/approval_flow_routes/{FREEE_APPROVAL_FLOW_ROUTE_ID}",
                headers=freee_headers("admin"),
                params={"company_id": FREEE_COMPANY_ID},
                timeout=30,
            )
            if not resp.ok:
                say(f":warning: 経路定義取得失敗: {resp.status_code} {resp.text}")
            else:
                say(f":mag: 経路定義: {json.dumps(resp.json(), ensure_ascii=False)[:3000]}")
        except Exception as e:
            say(f":warning: 経路定義取得失敗: {e}")
        return

    # 動作確認用コマンド: 事業所のメンバー一覧(freeeユーザーID込み)を取得する。
    # KNOWN_FREEE_USERS(Slackユーザー→freeeユーザーIDの対応表)に新しい人を
    # 追加する際、本人にfreee上のユーザーIDを確認してもらわなくても
    # ここから調べられるかを確認するためのもの。
    if text == "!debug_users":
        try:
            resp = requests.get(
                f"{FREEE_API_BASE}/api/1/users",
                headers=freee_headers("admin"),
                params={"company_id": FREEE_COMPANY_ID},
                timeout=30,
            )
            if not resp.ok:
                say(f":warning: メンバー一覧取得失敗: {resp.status_code} {resp.text}")
            else:
                say(f":mag: メンバー一覧: {json.dumps(resp.json(), ensure_ascii=False)[:3000]}")
        except Exception as e:
            say(f":warning: メンバー一覧取得失敗: {e}")
        return

    # 動作確認用コマンド: 現在有効な承認経路の一覧を取得する。
    # FREEE_APPROVAL_FLOW_ROUTE_ID(1431338)が「存在しないか既に削除された」
    # ことが!debug_routeで判明したため、現在使える経路IDを特定するために使う。
    if text == "!debug_routes_list":
        try:
            resp = requests.get(
                f"{FREEE_API_BASE}/api/1/approval_flow_routes",
                headers=freee_headers("admin"),
                params={"company_id": FREEE_COMPANY_ID},
                timeout=30,
            )
            if not resp.ok:
                say(f":warning: 経路一覧取得失敗: {resp.status_code} {resp.text}")
            else:
                say(f":mag: 経路一覧: {json.dumps(resp.json(), ensure_ascii=False)[:3000]}")
        except Exception as e:
            say(f":warning: 経路一覧取得失敗: {e}")
        return

    # ファイル無しの通常メッセージ(動作確認用)
    say("イベント受信成功")


@app.event("reaction_added")
def handle_reaction_added(event, say, logger):
    """確認メッセージへのリアクションでfreee申請を実行する"""
    logger.info(f"========== REACTION EVENT RECEIVED: {event} ==========")

    reaction = event.get("reaction")
    if reaction in APPROVE_REACTIONS:
        as_draft = False
    elif reaction in DRAFT_REACTIONS:
        as_draft = True
    else:
        logger.info(f"[reaction] 対象外のリアクションのため無視: {reaction}")
        return

    item = event.get("item", {})
    message_ts = item.get("ts")

    # 指定チャンネル以外は無視(テスト/本番の切り分け)
    if item.get("channel") != CONTRACT_CHANNEL_ID:
        logger.info(f"[reaction] 対象外チャンネルのため無視: {item.get('channel')}")
        return

    logger.info(f"[reaction] 現在のpending件数: {len(_pending_nda)}, 検索対象message_ts: {message_ts}")
    logger.info(f"[reaction] pending一覧: { {k: v.get('confirm_message_ts') for k, v in _pending_nda.items()} }")

    target_thread_ts = None
    for thread_ts, pending in _pending_nda.items():
        if pending.get("stage") == "awaiting_confirm" and pending.get("confirm_message_ts") == message_ts:
            target_thread_ts = thread_ts
            break

    if target_thread_ts is None:
        logger.warning(f"[reaction] 対応するpendingが見つかりません(message_ts={message_ts})。"
                        f"再デプロイ等でメモリ上の状態が失われた可能性があります。")
        return

    pending = _pending_nda.pop(target_thread_ts)

    try:
        # 投稿者本人のfreeeトークンが登録されていれば、それを使って
        # (=その人自身の名義で)証憑アップロード・申請作成を行う。
        # 未登録の場合は金子さんのトークンにフォールバックする。
        identity = resolve_freee_identity(pending.get("posting_slack_user_id"))
        logger.info(
            f"[freee] 使用するアイデンティティ: {identity} "
            f"(投稿者Slack ID: {pending.get('posting_slack_user_id')})"
        )
        logger.info(f"[freee] uploading file: {pending['filename']}")
        receipt_id = upload_file_to_freee(pending["raw_bytes"], pending["filename"], identity=identity)

        result = pending["gemini_result"]
        logger.info(f"[freee] creating approval request: {pending['filename']}")
        # タイトルは契約書に実際に印字されている表題を優先し、
        # 読み取れなければ書類種別、それも無ければファイル名にフォールバックする。
        contract_title = (
            (result.get("document_title") or "").strip()
            or (result.get("document_type") or "").strip()
            or os.path.splitext(pending["filename"])[0]
        )
        # 相手方は、既存のfreee取引先に一致するものがあればそのID(partner欄)、
        # 無ければ契約書に記載された正式名称を新規取引先名(single_line欄)として使う。
        counterparty_name = resolve_counterparty_name(result)
        partner_value, new_partner_name = resolve_counterparty_fields(counterparty_name)
        logger.info(
            f"[freee] タイトル: {contract_title!r} / 相手方名: {counterparty_name!r} / "
            f"partner欄: {partner_value!r} / 新規取引先欄: {new_partner_name!r}"
        )
        approval = create_nda_approval_request(
            title=contract_title,
            contract_date=result.get("contract_date") or datetime.date.today().isoformat(),
            receipt_id=receipt_id,
            section_id=pending["section_id"],
            method=pending["method"],
            approver_id=pending["approver_id"],
            applicant_id=pending.get("applicant_id"),
            partner_value=partner_value,
            new_partner_name=new_partner_name,
            mail_address=pending.get("mail_address", ""),
            draft=as_draft,
            identity=identity,
        )
        logger.info(
            f"[freee] 登録完了(draft={as_draft}, identity={identity})。"
            f"レスポンスのapplicant_id: {approval.get('applicant_id')}"
        )
        identity_note = (
            ""
            if identity == "user"
            else f"(申請者名義: {identity}のトークン)"
        )
        if as_draft:
            say(
                f":memo: freeeに下書きを保存しました(下書きID: {approval.get('id')}){identity_note}。\n"
                f"内容を確認・修正し、承認者を選んでfreee上で「申請」を押してください。",
                thread_ts=target_thread_ts,
            )
        else:
            say(
                f":white_check_mark: freeeへNDA契約締結申請を行いました"
                f"(申請番号: {approval.get('application_number')} / 承認者: {pending['approver_name']}){identity_note}",
                thread_ts=target_thread_ts,
            )
    except Exception as e:
        logger.exception("[freee] approval request creation failed")
        say(f":warning: freee申請の作成に失敗しました: {e}", thread_ts=target_thread_ts)


if __name__ == "__main__":
    print("BOT START")

    handler = SocketModeHandler(
        app,
        os.environ["SLACK_APP_TOKEN"],
    )
    handler.start()
